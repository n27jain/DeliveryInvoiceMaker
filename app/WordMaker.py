from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt
from datetime import datetime
from .Item import Item
from .Header import Header
import math

from .config import config      

class WordMaker:
    def __init__(self, clientName, fileName, itemList, taxPercent, header):
        self.taxPercent = taxPercent
        self.clientName = clientName
        self.itemList = itemList
        self.fileName = fileName
        self.time = datetime.today()
        self.dateString = ""+ str( self.time.month) + "/"+ str( self.time.day) + "/" + str( self.time.year)

        #these vars are used to create the header
        self.companyName = header.companyName if header.companyName else config.companyName
        self.address = header.companyAddress  if header.companyAddress else config.companyAddress

        tel = header.companyTel if header.companyTel else config.companyTel
        fax = header.companyFax  if header.companyFax else config.companyFax
        email = header.companyEmail if header.companyEmail else config.companyEmail
        
        self.tel = "Tel: " + str(tel)
        self.fax = "Fax: " + str(fax)
        self.email = "E-mail: " + str(email)

        self.note = header.companyNote if header.companyNote else config.companyNote

        self.document = Document()
        self.section = self.document.sections[0]

        self.makeHeader()
        self.makeData()
        self.makeFooter()
    def currencyFormat(self, value):
        value = float(value)
        return "${:,.2f}".format(value)
        
    def makeHeader(self):
        header = self.section.header
        l1 = header.add_paragraph(self.companyName)
        l1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        l1.paragraph_format.space_before = Pt(0)
        l1.paragraph_format.space_after = Pt(0)

        l2 = header.add_paragraph(self.address)
        l2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        l2.paragraph_format.space_before = Pt(0)
        l2.paragraph_format.space_after = Pt(0)

        l3 = header.add_paragraph(self.tel)
        l3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        l3.paragraph_format.space_before = Pt(0)
        l3.paragraph_format.space_after = Pt(0)

        l4 = header.add_paragraph(self.fax)
        l4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        l4.paragraph_format.space_before = Pt(0)
        l4.paragraph_format.space_after = Pt(0)

        l5 = header.add_paragraph(self.email)
        l5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        l5.paragraph_format.space_before = Pt(0)
        l5.paragraph_format.space_after = Pt(0)

    def makeData(self):
        self.document.add_paragraph("To: ")
        self.document.add_paragraph(self.clientName)
        self.document.add_paragraph()
        self.document.add_paragraph("Date: " + self.dateString)
        self.makeTable()
        self.document.add_paragraph()
        self.document.add_paragraph()

        self.document.add_paragraph(self.note)

    def makeFooter(self):
        footer = self.section.footer
        l1 = footer.add_paragraph(self.address)
        l1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        l1.paragraph_format.space_before = Pt(0)
        l1.paragraph_format.space_after = Pt(0)
        l2 = footer.add_paragraph(str(self.tel) + " " + str(self.fax) + " " + str(self.email))
        l2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        l2.paragraph_format.space_before = Pt(0)
        l2.paragraph_format.space_after = Pt(0)

    

    def makeTable(self):
        rowCount = len(self.itemList) + 5
        colCount = 4

        table = self.document.add_table(rows = rowCount, cols = colCount)
        table.style = 'Table Grid'
        
        table.rows[0].height = Inches(0.5)

        row = table.rows[0]
        cell = row.cells[0]
        cell.text = "Product Name"
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        row = table.rows[0]
        cell = row.cells[1]
        cell.text = "Price"
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        row = table.rows[0]
        cell = row.cells[2]
        cell.text = "Quantity"
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        row = table.rows[0]
        cell = row.cells[3]
        cell.text = "Amount"
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        subtotal = 0
        i = 1
        for item in self.itemList:
            amount = item.previousSalePrice[-1] * item.quantity
            subtotal += amount
            cell = table.cell(i,0)
            cell.text = str(item.name)
            cell = table.cell(i,1)
            cell.text = str(self.currencyFormat(item.previousSalePrice[-1]))
            cell = table.cell(i,2)
            cell.text = str(item.quantity)
           
            cell = table.cell(i,3)
            cell.text = str(self.currencyFormat(amount))
            i += 1

        hst = (0.13 * subtotal)
        total = (hst + subtotal)
        hst = round(hst, 2) 
        total = round(total, 2)
     
        row = table.rows[i+1]
        cell = row.cells[0]
        cell.text = "Sub Total"
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 2
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        row = table.rows[i+1]
        cell = row.cells[3]

        cell.text = self.currencyFormat(str(subtotal))

        row = table.rows[i+2]
        cell = row.cells[0]
        cell.text = "Hst 13%"
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 2
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        row = table.rows[i+2]
        cell = row.cells[3]
        cell.text = self.currencyFormat(str(hst))

        row = table.rows[i+3]
        cell = row.cells[0]
        cell.text = "Total"
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 2
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        row = table.rows[i+3]
        cell = row.cells[3]
        cell.text = self.currencyFormat(str(total))
    
    def save(self):
        self.document.save("app/"+ self.fileName + ".docx")
        return ("app/" + self.fileName + ".docx")



