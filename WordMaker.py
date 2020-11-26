from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt
from datetime import datetime
import math

class Item:
    def __init__(self, key, name, costPrice, previousSalePrice , notes, quantity):
        self.key = key
        self.name = name
        self.costPrice = costPrice
        self.previousSalePrice = previousSalePrice
        self.notes = notes
        self.quantity = quantity
        

class WordMaker:
    def __init__(self, clientName, fileName, itemList, taxPercent):
        self.taxPercent = taxPercent
        self.clientName = clientName
        self.itemList = itemList
        self.fileName = fileName
        self.time = datetime.today()
        self.dateString = ""+ str( self.time.month) + "/"+ str( self.time.day) + "/" + str( self.time.year)

        self.companyName = "2293984 ONTARIO INC."
        self.address = "147 Clarence Street, Unit#30, Brampton, ON. L6W 1T2"
        self.tel = "Tel: 905-497-6500"
        self.fax = "Fax: 905-497-5600"
        self.email = "E-mail: nathanpharmacy@gmail.com"
        self.note = "Due to the ongoing Covid-19 crisis, supply chains, shipping times and raw materials have been drastically affected. All prices and shipping estimates and times are subject to change without prior notice. Because of the nature of our products, we have a no returns and refund policy."
        self.document = Document()
        self.section = self.document.sections[0]

        self.makeHeader()
        self.makeData()
        self.makeFooter()

    def makeHeader(self):
        header = self.section.header
        l1 = header.add_paragraph(self.companyName)
        l1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        l1.paragraph_format.space_before = Pt(0)
        l1.paragraph_format.space_after = Pt(0)

        l2 = header.add_paragraph(self.address)
        l2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        l2.paragraph_format.space_before = Pt(0)
        l2.paragraph_format.space_after = Pt(0)

        l3 = header.add_paragraph(self.tel)
        l3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        l3.paragraph_format.space_before = Pt(0)
        l3.paragraph_format.space_after = Pt(0)

        l4 = header.add_paragraph(self.fax)
        l4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        l4.paragraph_format.space_before = Pt(0)
        l4.paragraph_format.space_after = Pt(0)

        l5 = header.add_paragraph(self.email)
        l5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        l5.paragraph_format.space_before = Pt(0)
        l5.paragraph_format.space_after = Pt(0)

    def makeData(self):
        self.document.add_paragraph("To: ")
        self.document.add_paragraph(self.clientName)
        self.document.add_paragraph()
        self.document.add_paragraph("Date: " + self.dateString)
        self.makeTable()
        self.document.add_paragraph()
        self.document.add_paragraph()

        self.document.add_paragraph(self.note)

    def makeFooter(self):
        footer = self.section.footer
        l1 = footer.add_paragraph(self.address)
        l1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        l1.paragraph_format.space_before = Pt(0)
        l1.paragraph_format.space_after = Pt(0)
        l2 = footer.add_paragraph(str(self.tel) + " " + str(self.fax) + " " + str(self.email))
        l2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        l2.paragraph_format.space_before = Pt(0)
        l2.paragraph_format.space_after = Pt(0)

    def roundup(self, var):
        return (math.ceil(var / 100.0)) * 100

    def makeTable(self):
        rowCount = len(self.itemList) + 5
        colCount = 4

        table = self.document.add_table(rows = rowCount, cols = colCount)
        table.style = 'Table Grid'
        
        table.rows[0].height = Inches(0.5)

        row = table.rows[0]
        cell = row.cells[0]
        cell.text = "Product Name"
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        row = table.rows[0]
        cell = row.cells[1]
        cell.text = "Price"
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        row = table.rows[0]
        cell = row.cells[2]
        cell.text = "Quantity"
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        row = table.rows[0]
        cell = row.cells[3]
        cell.text = "Amount"
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        subtotal = 0
        i = 1
        for item in self.itemList:
            amount = item.previousSalePrice * item.quantity
            subtotal += amount
            cell = table.cell(i,0)
            cell.text = str(item.name)
            cell = table.cell(i,1)
            cell.text = str(item.previousSalePrice)
            cell = table.cell(i,2)
            cell.text = str(item.quantity)
           
            cell = table.cell(i,3)
            cell.text = str(amount)
            i += 1

        hst = (0.13 * subtotal)
        total = (hst + subtotal)
        hst = round(hst, 2) 
        total = round(total, 2)
     
        row = table.rows[i+1]
        cell = row.cells[0]
        cell.text = "Sub Total"
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 2
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        row = table.rows[i+1]
        cell = row.cells[3]

        cell.text = str(subtotal)

        row = table.rows[i+2]
        cell = row.cells[0]
        cell.text = "Hst 13%"
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 2
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        row = table.rows[i+2]
        cell = row.cells[3]
        cell.text = str(hst)

        row = table.rows[i+3]
        cell = row.cells[0]
        cell.text = "Total"
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 2
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True

        row = table.rows[i+3]
        cell = row.cells[3]
        cell.text = str(total)
    
    def save(self):
        self.document.save(self.fileName + ".docx")
        return (self.fileName + ".docx")



